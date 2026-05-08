"""Document parsers for .docx and .pdf files."""

import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def parse_document(file_path: str) -> dict:
    """Parse a .docx or .pdf file and return structured content.

    Returns a dict with keys:
        file_type, text, paragraphs, headings, table_count, tables,
        image_count, page_count, is_page_count_estimated
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    if path.stat().st_size == 0:
        raise ValueError(f"File is empty: {file_path}")

    ext = path.suffix.lower()
    if ext == ".docx":
        return _parse_docx(file_path)
    elif ext == ".pdf":
        return _parse_pdf(file_path)
    else:
        raise ValueError(f"Unsupported file type '{ext}'. Supported: .docx, .pdf")


# -- DOCX parser --------------------------------------------------------------

def _parse_docx(file_path: str) -> dict:
    """Parse a Word (.docx) document using python-docx."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is not installed. Run: pip install python-docx")

    doc = Document(file_path)
    headings: list[dict] = []
    paragraphs: list[str] = []
    all_text_parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        all_text_parts.append(text)
        style = para.style.name

        if style.startswith("Heading") or style == "Title":
            level = 0 if style == "Title" else _heading_level(style)
            headings.append({"text": text, "level": level, "style": style})
        else:
            paragraphs.append(text)

    # Count images via relationships
    image_count = sum(
        1 for rel in doc.part.rels.values() if "image" in rel.target_ref
    )

    # Extract tables
    table_count = len(doc.tables)
    tables: list[list[list[str]]] = []
    for table in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        tables.append(rows)

    full_text = "\n\n".join(all_text_parts)
    word_count = len(full_text.split())
    estimated_pages = max(1, round(word_count / 300))

    return {
        "file_type": "docx",
        "text": full_text,
        "paragraphs": paragraphs,
        "headings": headings,
        "table_count": table_count,
        "tables": tables,
        "image_count": image_count,
        "page_count": estimated_pages,
        "is_page_count_estimated": True,
    }


# -- PDF parser ----------------------------------------------------------------

def _parse_pdf(file_path: str) -> dict:
    """Parse a PDF document using pdfplumber."""
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("pdfplumber is not installed. Run: pip install pdfplumber")

    paragraphs: list[str] = []
    all_text_parts: list[str] = []
    headings: list[dict] = []
    page_count = 0

    with pdfplumber.open(file_path) as pdf:
        page_count = len(pdf.pages)
        body_font_size = _detect_body_font_size(pdf)

        for page in pdf.pages:
            text = page.extract_text(x_tolerance=3, y_tolerance=3)
            if not text:
                continue

            # Split raw text into paragraph blocks
            raw_blocks = [b.strip() for b in text.split("\n\n") if b.strip()]
            if not raw_blocks:
                raw_blocks = [ln.strip() for ln in text.splitlines() if ln.strip()]

            for block in raw_blocks:
                all_text_parts.append(block)

            # Detect headings via font size on this page
            page_headings = _extract_pdf_headings(page, body_font_size)
            headings.extend(page_headings)

            # Paragraphs = blocks that are NOT headings
            heading_texts = {h["text"] for h in page_headings}
            for block in raw_blocks:
                if block not in heading_texts:
                    paragraphs.append(block)

    full_text = "\n\n".join(all_text_parts)

    return {
        "file_type": "pdf",
        "text": full_text,
        "paragraphs": paragraphs,
        "headings": headings,
        "table_count": 0,
        "tables": [],
        "image_count": 0,
        "page_count": page_count,
        "is_page_count_estimated": False,
    }


# -- Helpers -------------------------------------------------------------------

def _heading_level(style_name: str) -> int:
    """Extract the numeric heading level from a Word style name like 'Heading 2'."""
    parts = style_name.split()
    return int(parts[-1]) if parts and parts[-1].isdigit() else 1


def _detect_body_font_size(pdf) -> float:
    """Return the most common font size (proxy for body text)."""
    size_freq: dict[float, int] = {}
    sample_pages = pdf.pages[: min(5, len(pdf.pages))]
    for page in sample_pages:
        for char in page.chars:
            size = round(float(char.get("size", 0)), 1)
            if size > 0:
                size_freq[size] = size_freq.get(size, 0) + 1

    if not size_freq:
        return 12.0
    return max(size_freq, key=lambda s: size_freq[s])


def _extract_pdf_headings(page, body_font_size: float) -> list[dict]:
    """Detect heading lines on a single PDF page via font-size heuristics."""
    if not page.chars:
        return []

    # Group characters by vertical position (same line)
    lines_by_top: dict[int, list] = {}
    for char in page.chars:
        top = round(float(char.get("top", 0)))
        lines_by_top.setdefault(top, []).append(char)

    headings: list[dict] = []
    seen: set[str] = set()

    for top in sorted(lines_by_top):
        chars_in_line = lines_by_top[top]
        sizes = [
            round(float(c.get("size", 0)), 1)
            for c in chars_in_line
            if c.get("size", 0) > 0
        ]
        if not sizes:
            continue

        avg_size = sum(sizes) / len(sizes)
        if avg_size < body_font_size * 1.15:
            continue

        line_text = "".join(c.get("text", "") for c in chars_in_line).strip()
        if not line_text or len(line_text) < 2 or line_text in seen:
            continue

        seen.add(line_text)
        level = 1 if avg_size >= body_font_size * 1.4 else 2
        headings.append({
            "text": line_text,
            "level": level,
            "style": f"PDF-H{level} ({avg_size:.0f}pt)",
        })

    return headings
